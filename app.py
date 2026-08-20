import streamlit as st
import os
import pypdf
import docx
from io import BytesIO
import sqlite3
import re
import requests
import asyncio
import edge_tts
import tempfile
from PIL import Image
from youtube_transcript_api import YouTubeTranscriptApi
import json
import urllib.parse
from fpdf import FPDF
import uuid
from streamlit_cookies_controller import CookieController

# 1. Page Configuration
st.set_page_config(page_title=" Vishnu's SSLC AI Master 📚", page_icon="🎓", layout="wide", initial_sidebar_state="expanded")
# ==========================================
# 🔑 FREE GROQ API CONFIGURATION
# ==========================================
GROQ_API_KEY = "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"
GROQ_URL = "https://api.groq.com/openai/v1/chat/completions"
DEFAULT_MODEL = "openai/gpt-oss-120b"  # Free & Fast Model

def call_aimlapi(messages, model=DEFAULT_MODEL, temperature=0.7):
    headers = {
        "Authorization": f"Bearer {GROQ_API_KEY}",
        "Content-Type": "application/json"
    }
    payload = {
        "model": model,
        "messages": messages,
        "temperature": temperature
    }
    try:
        response = requests.post(GROQ_URL, headers=headers, json=payload, timeout=60)
        if response.status_code == 200:
            res_json = response.json()
            return res_json["choices"][0]["message"]["content"]
        else:
            return f"API Error ({response.status_code}): {response.text}"
    except Exception as e:
        return f"Request Error: {str(e)}"


# 3. Cookie Controller for Persistent Login
controller = CookieController()

# --- Database Setup for Chat History ---
def init_db():
    conn = sqlite3.connect('chat_history.db', check_same_thread=False)
    c = conn.cursor()
    c.execute('''CREATE TABLE IF NOT EXISTS chat_sessions (session_id TEXT PRIMARY KEY, email TEXT, title TEXT)''')
    c.execute('''CREATE TABLE IF NOT EXISTS chat_messages (id INTEGER PRIMARY KEY AUTOINCREMENT, session_id TEXT, role TEXT, content TEXT)''')
    conn.commit()
    return conn

conn = init_db()

def save_chat_message(session_id, email, role, content):
    c = conn.cursor()
    c.execute("SELECT session_id FROM chat_sessions WHERE session_id=?", (session_id,))
    if not c.fetchone():
        title = content[:25] + "..." if len(content) > 25 else content
        c.execute("INSERT INTO chat_sessions (session_id, email, title) VALUES (?, ?, ?)", (session_id, email, title))
    c.execute("INSERT INTO chat_messages (session_id, role, content) VALUES (?, ?, ?)", (session_id, role, content))
    conn.commit()

def get_chat_sessions(email):
    c = conn.cursor()
    c.execute("SELECT session_id, title FROM chat_sessions WHERE email = ? ORDER BY rowid DESC", (email,))
    return c.fetchall()

def get_chat_messages(session_id):
    c = conn.cursor()
    c.execute("SELECT role, content FROM chat_messages WHERE session_id = ? ORDER BY id ASC", (session_id,))
    return c.fetchall()

# --- Session States ---
if "logged_in" not in st.session_state: 
    st.session_state.logged_in = False
    st.session_state.user_email = ""

if "current_session_id" not in st.session_state: 
    st.session_state.current_session_id = str(uuid.uuid4())
    
if "selected_voice" not in st.session_state:
    st.session_state.selected_voice = "ml-IN-SobhanaNeural"

if "app_mode_select" not in st.session_state:
    st.session_state.app_mode_select = "6. യഥാർത്ഥ ചാറ്റ് (Chatbot UI)"

# Persistent Login Check
if not st.session_state.logged_in:
    cookie_email = controller.get('user_email')
    if cookie_email:
        st.session_state.logged_in = True
        st.session_state.user_email = cookie_email

for i in range(8):
    if f"out_{i}" not in st.session_state:
        st.session_state[f"out_{i}"] = ""

# 4. Voices Dictionary
VOICES = {
    "മലയാളം - ശോഭന (Female)": "ml-IN-SobhanaNeural",
    "മലയാളം - മിഥുൻ (Male)": "ml-IN-MidhunNeural",
    "English (India) - Neerja (Female)": "en-IN-NeerjaNeural",
    "English (India) - Prabhat (Male)": "en-IN-PrabhatNeural"
}

# --- Helper Functions ---
def create_docx(text):
    doc = docx.Document()
    doc.add_heading('SSLC AI Smart Notes', 0)
    doc.add_paragraph(text)
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

def create_pdf(text):
    try:
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.set_font("Helvetica", size=12)
        safe_text = text.encode('latin-1', 'replace').decode('latin-1')
        pdf.multi_cell(0, 10, txt=safe_text)
        return bytes(pdf.output())
    except Exception:
        return b""

def create_audio_improved(text, voice):
    async def _generate():
        communicate = edge_tts.Communicate(text, voice)
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".mp3")
        temp_path = temp_file.name
        temp_file.close()
        await communicate.save(temp_path)
        return temp_path
    try:
        try:
            loop = asyncio.get_event_loop()
        except RuntimeError:
            loop = asyncio.new_event_loop()
            asyncio.set_event_loop(loop)
            
        if loop.is_running():
            import nest_asyncio
            nest_asyncio.apply()
            audio_file_path = loop.run_until_complete(_generate())
        else:
            audio_file_path = asyncio.run(_generate())

        with open(audio_file_path, "rb") as f:
            audio_bytes = f.read()
        os.remove(audio_file_path)
        return audio_bytes
    except Exception:
        return None

def extract_json(raw_text):
    match = re.search(r'\[.*\]|\{.*\}', raw_text, re.DOTALL)
    if match:
        return json.loads(match.group(0))
    return json.loads(raw_text)

def render_smart_actions(text_content, feature_id):
    st.markdown("---")
    st.markdown("### 🛠️ സേവ് ചെയ്യാം (Save & Share)")
    col1, col2, col3, col4 = st.columns(4)
    with col1: 
        st.download_button("📄 Word (Docx)", create_docx(text_content), f"Result_{feature_id}.docx", key=f"docx_{feature_id}")
    with col2: 
        st.download_button("📕 PDF", create_pdf(text_content), f"Result_{feature_id}.pdf", key=f"pdf_{feature_id}")
    with col3:
        if st.button("🎧 ഓഡിയോ കേൾക്കാം", key=f"audio_{feature_id}"):
            with st.spinner("ഓഡിയോ തയ്യാറാക്കുന്നു..."):
                audio_bytes = create_audio_improved(text_content, st.session_state.selected_voice)
                if audio_bytes: 
                    st.audio(audio_bytes, format='audio/mp3')
                else: 
                    st.error("ഓഡിയോ ഉണ്ടാക്കുന്നതിൽ തകരാർ.")
    with col4:
        with st.expander("📝 കോപ്പി ചെയ്യാൻ"): 
            st.code(text_content, language='text')

def extract_text_from_pdf(pdf_file):
    reader = pypdf.PdfReader(pdf_file)
    return "".join([page.extract_text() or "" for page in reader.pages])

def download_gdrive_pdf(url):
    try:
        file_id_match = re.search(r'[/=]([-\w]{25,})', url)
        if not file_id_match: 
            return None, "ലിങ്കിൽ നിന്നും ഫയൽ ഐഡി കണ്ടെത്താൻ കഴിഞ്ഞില്ല."
        download_url = f"https://drive.google.com/uc?export=download&id={file_id_match.group(1)}"
        response = requests.get(download_url, allow_redirects=True)
        if response.status_code == 200 and response.content.startswith(b'%PDF'):
            return BytesIO(response.content), "Success"
        else: 
            return None, "ഡൗൺലോഡ് ചെയ്തത് PDF അല്ല."
    except Exception as e: 
        return None, f"Error: {str(e)}"

# --- Custom Styling ---
st.markdown("""
    <style>
    .stApp { background-color: #0b0f19 !important; color: #ffffff !important; }
    [data-testid="stSidebar"] { background-color: #060911 !important; border-right: 1px solid #1f2937; }
    .stChatMessage { background-color: transparent !important; }
    [data-testid="chatAvatarIcon-user"] { background-color: #6366f1 !important; }
    [data-testid="chatAvatarIcon-assistant"] { background-color: #1e293b !important; }
    .stChatMessage div[data-testid="stMarkdownContainer"] {
        background-color: #111827;
        padding: 15px;
        border-radius: 12px;
        border: 1px solid #1f2937;
    }
    .stButton>button { 
        background: linear-gradient(135deg, #4f46e5, #7c3aed) !important; 
        color: #FFFFFF !important; 
        border: none !important; 
        border-radius: 12px !important; 
        font-weight: 600 !important;
    }
    [data-testid="stChatInput"] { background-color: #111827 !important; border-radius: 15px !important; border: 1px solid #374151 !important; }
    </style>
""", unsafe_allow_html=True)

# --- Sidebar Setup ---
st.sidebar.markdown("## 🔮 NEXORA AI\n<span style='color:gray; font-size: 14px;'>AI Companion</span>", unsafe_allow_html=True)

def start_new_chat():
    st.session_state.current_session_id = str(uuid.uuid4())
    st.session_state.app_mode_select = "6. യഥാർത്ഥ ചാറ്റ് (Chatbot UI)"

def load_existing_chat(sess_id):
    st.session_state.current_session_id = sess_id
    st.session_state.app_mode_select = "6. യഥാർത്ഥ ചാറ്റ് (Chatbot UI)"

st.sidebar.button("➕ New Chat", use_container_width=True, on_click=start_new_chat)

app_mode = st.sidebar.radio(
    "Explore Options:",
    ["1. യഥാർത്ഥ ചാറ്റ് (Chatbot UI)",
     "2. പാഠപുസ്തകം / നോട്ട്സ് (PDF Chat)",
     "3. ചിത്രങ്ങൾ നൽകി പഠിക്കാം (Image Analysis)",
     "4. AI Mock Test (ക്വിസ്)",
     "5. വോയിസ് ഇൻപുട്ട് (സംസാരിച്ച് ചോദിക്കാം)",
     "6. YouTube Video Summarizer",
     "7. സ്റ്റഡി പ്ലാനർ (Study Planner)",
     "8. ഫ്ലാഷ് കാർഡുകൾ (Quick Revision)"],
    key="app_mode_select"
)

st.sidebar.markdown("---")
st.sidebar.markdown("### Recent Chats")

current_email = st.session_state.user_email if st.session_state.logged_in else f"guest_{st.session_state.current_session_id}"

sessions = get_chat_sessions(current_email)
if sessions:
    for sess_id, title in sessions:
        st.sidebar.button(f"💬 {title}", key=f"btn_{sess_id}", use_container_width=True, on_click=load_existing_chat, args=(sess_id,))

st.sidebar.markdown("---")
if not st.session_state.logged_in:
    login_email = st.sidebar.text_input("📧 Email Address (For Login):")
    if st.sidebar.button("Login", use_container_width=True) and login_email:
        controller.set('user_email', login_email)
        st.session_state.logged_in = True
        st.session_state.user_email = login_email
        st.rerun()
else:
    st.sidebar.success(f"👤 {st.session_state.user_email}")
    if st.sidebar.button("Logout", use_container_width=True):
        controller.remove('user_email')
        st.session_state.logged_in = False
        st.session_state.user_email = ""
        st.session_state.current_session_id = str(uuid.uuid4())
        st.rerun()

st.sidebar.title("🔊 Audio Settings")
selected_voice_name = st.sidebar.selectbox("വോയിസ് തിരഞ്ഞെടുക്കുക:", list(VOICES.keys()))
st.session_state.selected_voice = VOICES[selected_voice_name]

# =============================================================================
# 1. Chatbot UI
# =============================================================================
if app_mode == "1. യഥാർത്ഥ ചാറ്റ് (Chatbot UI)":
    st.markdown(f"<h2>Hello, {current_email.split('@')[0]}! 👋</h2>", unsafe_allow_html=True)
    st.caption("How can I help you today?")
    
    session_id = st.session_state.current_session_id
    db_messages = get_chat_messages(session_id)
    
    for role, content in db_messages:
        with st.chat_message(role): 
            st.markdown(content)
        
    with st.popover("➕ അറ്റാച്ച് ചെയ്യുക"):
        st.write("ഫയലുകൾ അപ്‌ലോഡ് ചെയ്യുക:")
        chat_upload_method = st.radio("രീതി:", ["PDF File", "Google Drive PDF Link"])
        chat_pdf_text = ""
        
        if chat_upload_method == "PDF File":
            chat_pdfs = st.file_uploader("PDF:", type=["pdf"], accept_multiple_files=True)
            if chat_pdfs: 
                chat_pdf_text = "".join([extract_text_from_pdf(p) for p in chat_pdfs])
        elif chat_upload_method == "Google Drive PDF Link":
            chat_url = st.text_input("Drive Link:")
            if chat_url:
                with st.spinner("Downloading..."):
                    pdf_bytes, _ = download_gdrive_pdf(chat_url)
                    if pdf_bytes: 
                        chat_pdf_text = extract_text_from_pdf(pdf_bytes)

    if prompt := st.chat_input("നിങ്ങളുടെ സംശയം ചോദിക്കുക..."):
        with st.chat_message("user"): 
            st.markdown(prompt)
        
        display_prompt = prompt
        if chat_pdf_text: display_prompt += " [PDF Attached]"
        save_chat_message(session_id, current_email, "user", display_prompt)

        with st.spinner("ചിന്തിക്കുന്നു..."):
            messages = [{"role": "system", "content": "You are NEXORA AI, an expert Kerala SSLC educational assistant. Answer questions clearly in Malayalam unless requested in English."}]
            for r, c in db_messages:
                messages.append({"role": r, "content": c})
            
            full_prompt = prompt
            if chat_pdf_text:
                full_prompt += f"\n\nContext from PDF: {chat_pdf_text[:8000]}"
            messages.append({"role": "user", "content": full_prompt})
            
            ai_reply = call_aimlapi(messages)
            st.session_state.out_5 = ai_reply
            with st.chat_message("assistant"): 
                st.markdown(ai_reply)
            save_chat_message(session_id, current_email, "assistant", ai_reply)
            
    if st.session_state.out_5:
        render_smart_actions(st.session_state.out_5, "5")

# =============================================================================
# 2. PDF / Notes Chat
# =============================================================================
elif app_mode == "2. പാഠപുസ്തകം / നോട്ട്സ് (PDF Chat)":
    st.header("📄 PDF നോട്ട്സ് അനാലിസിസ്")
    pdf_text = ""
    upload_method = st.radio("എങ്ങനെയാണ് നോട്ട്സ് നൽകുന്നത്?", ("വേണ്ട (ചോദ്യം മാത്രം)", "PDF അപ്‌ലോഡ് ചെയ്യുക", "Google Drive ലിങ്ക് നൽകുക"))
    
    if upload_method == "PDF അപ്‌ലോഡ് ചെയ്യുക":
        uploaded_pdfs = st.file_uploader("📄 പാഠപുസ്തകം (PDF):", type=["pdf"], accept_multiple_files=True)
        if uploaded_pdfs: 
            pdf_text = "".join([extract_text_from_pdf(p) for p in uploaded_pdfs])
    elif upload_method == "Google Drive ലിങ്ക് നൽകുക":
        gdrive_url = st.text_input("🔗 ഗൂഗിൾ ഡ്രൈവ് ലിങ്ക്:")
        if gdrive_url:
            pdf_bytes, status = download_gdrive_pdf(gdrive_url)
            if pdf_bytes: 
                pdf_text = extract_text_from_pdf(pdf_bytes)
                
    user_input = st.text_area("ചോദ്യം നൽകുക:")
    if st.button("ഉത്തരം കണ്ടെത്തുക 🚀") and (user_input or pdf_text):
        combined = user_input + (f"\n\nPDF Context:\n{pdf_text[:10000]}" if pdf_text else "")
        with st.spinner('AI ഉത്തരം തയ്യാറാക്കുന്നു...'):
            messages = [
                {"role": "system", "content": "Act as an expert Kerala SSLC teacher. Explain topics simply and accurately in Malayalam."},
                {"role": "user", "content": combined}
            ]
            ai_reply = call_aimlapi(messages)
            st.session_state.out_0 = ai_reply
            
            save_chat_message(st.session_state.current_session_id, current_email, "user", f"[PDF Analysis] {user_input}")
            save_chat_message(st.session_state.current_session_id, current_email, "assistant", ai_reply)
            
    if st.session_state.out_0:
        st.write(st.session_state.out_0)
        render_smart_actions(st.session_state.out_0, "0")

# =============================================================================
# 3. Image Analysis
# =============================================================================
elif app_mode == "3. ചിത്രങ്ങൾ നൽകി പഠിക്കാം (Image Analysis)":
    st.header("📷 ചിത്രങ്ങൾ നൽകി പഠിക്കാം")
    st.info("AIMLAPI text models process text. For visual questions, describe the diagram or problem below.")
    uploaded_files = st.file_uploader("ചിത്രങ്ങൾ അപ്‌ലോഡ് ചെയ്യുക (Preview)", type=["jpg", "png", "jpeg"], accept_multiple_files=True)
    
    if uploaded_files:
        cols = st.columns(4) 
        for idx, file in enumerate(uploaded_files[:20]):
            img = Image.open(file)
            with cols[idx % 4]: 
                st.image(img, use_container_width=True)
                
    prompt = st.text_input("ചിത്രത്തിലെ ചോദ്യം അല്ലെങ്കിൽ വിഷയം ഇവിടെ ടൈപ്പ് ചെയ്യുക:")
    if st.button("കണ്ടെത്തുക") and prompt:
        with st.spinner("പരിശോധിക്കുന്നു..."):
            messages = [
                {"role": "system", "content": "Act as an SSLC teacher. Explain the question or concept clearly in Malayalam."},
                {"role": "user", "content": prompt}
            ]
            ai_reply = call_aimlapi(messages)
            st.session_state.out_1 = ai_reply
            
            save_chat_message(st.session_state.current_session_id, current_email, "user", f"[Image/Topic Analysis] {prompt}")
            save_chat_message(st.session_state.current_session_id, current_email, "assistant", ai_reply)
            
    if st.session_state.out_1:
        st.write(st.session_state.out_1)
        render_smart_actions(st.session_state.out_1, "1")

# =============================================================================
# 4. Mock Test
# =============================================================================
elif app_mode == "4. AI Mock Test (ക്വിസ്)":
    st.header("📝 സ്വയം പരീക്ഷിക്കാം")
    topic = st.text_input("ഏത് വിഷയത്തിലാണ് ടെസ്റ്റ് വേണ്ടത്?")
    if "quiz_data" not in st.session_state: 
        st.session_state.quiz_data = None
    
    if st.button("ക്വിസ് തുടങ്ങുക") and topic:
        with st.spinner("ചോദ്യങ്ങൾ തയ്യാറാക്കുന്നു..."):
            try:
                prompt_text = (
                    f"Create a 5-question multiple choice quiz on '{topic}' for Kerala SSLC 10th-grade. "
                    f"Return strictly a JSON array of objects with keys: 'question' (Malayalam), 'options' (array of 4 options in Malayalam), and 'answer' (exact match with one option). "
                    f"Output raw JSON only without markdown code blocks."
                )
                messages = [
                    {"role": "system", "content": "You are a JSON-only API. You output raw valid JSON arrays."},
                    {"role": "user", "content": prompt_text}
                ]
                raw_response = call_aimlapi(messages, temperature=0.2)
                st.session_state.quiz_data = extract_json(raw_response)
                quiz_str = f"📝 Mock Test: {topic}\n\n"
                for i, q in enumerate(st.session_state.quiz_data): 
                    quiz_str += f"Q{i+1}: {q['question']}\nAns: {q['answer']}\n\n"
                st.session_state.out_2 = quiz_str
                
                save_chat_message(st.session_state.current_session_id, current_email, "user", f"[Mock Test] Generate quiz on {topic}")
                save_chat_message(st.session_state.current_session_id, current_email, "assistant", quiz_str)
            except Exception as e: 
                st.error(f"Error parsing quiz: {e}")
                
    if st.session_state.quiz_data:
        user_answers = {}
        for i, q in enumerate(st.session_state.quiz_data):
            user_answers[i] = st.radio(f"{i+1}. {q['question']}", q['options'], key=f"q_{i}")
        if st.button("പരിശോധിക്കുക"):
            score = sum([1 for i, q in enumerate(st.session_state.quiz_data) if user_answers.get(i) == q['answer']])
            st.header(f"സ്കോർ: {score}/5")
            st.success("ഉത്തരങ്ങൾ താഴെ സേവ് ചെയ്യാം!")
            
    if st.session_state.out_2:
        render_smart_actions(st.session_state.out_2, "2")

# =============================================================================
# 5. Voice Input
# =============================================================================
elif app_mode == "5. വോയിസ് ഇൻപുട്ട് (സംസാരിച്ച് ചോദിക്കാം)":
    st.header("🎙️ സംസാരിച്ച് ചോദ്യം ചോദിക്കാം")
    st.info("താഴെയുള്ള ബോക്സിൽ ചോദ്യം നൽകുക:")
    voice_query = st.text_input("ചോദ്യം ഇവിടെ നൽകുക:")
    
    if st.button("ഉത്തരം കണ്ടെത്തുക") and voice_query:
        with st.spinner("കണ്ടെത്തുന്നു..."):
            messages = [
                {"role": "system", "content": "Answer clearly and helpfully in Malayalam for an SSLC student."},
                {"role": "user", "content": voice_query}
            ]
            ai_reply = call_aimlapi(messages)
            st.session_state.out_3 = ai_reply
            
            save_chat_message(st.session_state.current_session_id, current_email, "user", f"[Voice/Text Input] {voice_query}")
            save_chat_message(st.session_state.current_session_id, current_email, "assistant", ai_reply)
            
    if st.session_state.out_3:
        st.write(st.session_state.out_3)
        render_smart_actions(st.session_state.out_3, "3")

# =============================================================================
# 6. YouTube Summarizer
# =============================================================================
elif app_mode == "6. YouTube Video Summarizer":
    st.header("📺 YouTube ക്ലാസ്സ് നോട്ട്സ്")
    yt_url = st.text_input("YouTube Link നൽകുക:")
    if st.button("Summary തയ്യാറാക്കുക") and yt_url:
        try:
            parsed_url = urllib.parse.urlparse(yt_url)
            video_id = parsed_url.path[1:] if parsed_url.hostname == 'youtu.be' else urllib.parse.parse_qs(parsed_url.query)['v'][0]
            with st.spinner("ശേഖരിക്കുന്നു..."):
                transcript = YouTubeTranscriptApi.get_transcript(video_id, languages=['en', 'ml'])
                full_transcription = " ".join([i['text'] for i in transcript])
                
                messages = [
                    {"role": "system", "content": "You summarize educational video transcripts into clear Malayalam bullet-point notes for SSLC students."},
                    {"role": "user", "content": f"Summarize in Malayalam:\n\n{full_transcription[:6000]}"}
                ]
                ai_reply = call_aimlapi(messages)
                st.session_state.out_4 = ai_reply
                
                save_chat_message(st.session_state.current_session_id, current_email, "user", f"[YouTube Summary] {yt_url}")
                save_chat_message(st.session_state.current_email, current_email, "assistant", ai_reply)
        except Exception: 
            st.error("സബ്ടൈറ്റിലുകൾ ലഭ്യമായില്ല അല്ലെങ്കിൽ ലിങ്ക് തെറ്റാണ്.")
        
    if st.session_state.out_4:
        st.write(st.session_state.out_4)
        render_smart_actions(st.session_state.out_4, "4")

# =============================================================================
# 7. Study Planner
# =============================================================================
elif app_mode == "7. സ്റ്റഡി പ്ലാനർ (Study Planner)":
    st.header("📅 സ്റ്റഡി പ്ലാനർ")
    days = st.number_input("ദിവസങ്ങൾ?", min_value=1, value=30)
    hours = st.number_input("മണിക്കൂർ?", min_value=1, value=3)
    subjects = st.text_area("വിഷയങ്ങൾ:", "Physics, Chemistry, Maths")
    if st.button("തയ്യാറാക്കുക"):
        with st.spinner("പ്ലാൻ തയ്യാറാക്കുന്നു..."):
            messages = [
                {"role": "system", "content": "You are an SSLC academic coach. Build clear markdown study timetables and explain revision strategies in Malayalam."},
                {"role": "user", "content": f"Create a study timetable. Days: {days}, Hours/day: {hours}, Subjects: {subjects}."}
            ]
            ai_reply = call_aimlapi(messages)
            st.session_state.out_6 = ai_reply
            
            save_chat_message(st.session_state.current_session_id, current_email, "user", f"[Study Planner] {days} Days, {hours} Hrs/Day, Subjects: {subjects}")
            save_chat_message(st.session_state.current_session_id, current_email, "assistant", ai_reply)
            
    if st.session_state.out_6:
        st.markdown(st.session_state.out_6)
        render_smart_actions(st.session_state.out_6, "6")

# =============================================================================
# 8. Flash Cards
# =============================================================================
elif app_mode == "8. ഫ്ലാഷ് കാർഡുകൾ (Quick Revision)":
    st.header("⚡ ഫ്ലാഷ് കാർഡുകൾ")
    flash_topic = st.text_input("ഏത് വിഷയമാണ് റിവിഷൻ ചെയ്യേണ്ടത്?")
    if "flash_cards_data" not in st.session_state: 
        st.session_state.flash_cards_data = None
    
    if st.button("തയ്യാറാക്കുക") and flash_topic:
        with st.spinner("ഉണ്ടാക്കുന്നു..."):
            try:
                flash_prompt = (
                    f"Create 10 key flashcards for '{flash_topic}' for 10th-grade SSLC. "
                    f"Provide explanation in English and Malayalam. "
                    f"Return strictly a JSON array of objects with keys 'title' and 'description'. Output raw JSON only."
                )
                messages = [
                    {"role": "system", "content": "You are a JSON-only API that outputs valid JSON arrays."},
                    {"role": "user", "content": flash_prompt}
                ]
                raw_response = call_aimlapi(messages, temperature=0.3)
                st.session_state.flash_cards_data = extract_json(raw_response)
                
                cards_str = f"⚡ ഫ്ലാഷ് കാർഡുകൾ: {flash_topic}\n\n"
                for card in st.session_state.flash_cards_data: 
                    cards_str += f"📌 {card['title']}\n{card['description']}\n\n"
                st.session_state.out_7 = cards_str
                
                save_chat_message(st.session_state.current_session_id, current_email, "user", f"[Flash Cards] Generate on {flash_topic}")
                save_chat_message(st.session_state.current_session_id, current_email, "assistant", cards_str)
            except Exception as e:
                st.error(f"Error parsing flash cards: {e}")
            
    if st.session_state.flash_cards_data:
        cols = st.columns(3)
        for idx, card in enumerate(st.session_state.flash_cards_data):
            with cols[idx % 3]:
                with st.container(border=True):
                    st.subheader(card.get('title', 'Concept'))
                    st.write(card.get('description', ''))
                    
    if st.session_state.out_7:
        render_smart_actions(st.session_state.out_7, "7")
